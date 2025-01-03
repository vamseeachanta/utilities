# Third party imports
import pypandoc
from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Add a 3x2 table
table = doc.add_table(rows=4, cols=3)

# Path to the image
image_path = 'docs\sub_ms_office\SD_Photo.jpeg'

# Insert the image into each cell and resize it
for row in table.rows:
    for cell in row.cells:
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(2), height=Inches(2))

# Save the document
word_file = 'docs\sub_ms_office\Passport_Photo_Grid.docx'
doc.save(word_file)

# Convert the Word file to PDF
try:
    pdf_file = 'docs\sub_ms_office\Passport_Photo_Grid.pdf'
    pypandoc.convert_file(word_file, 'pdf', outputfile=pdf_file)

    print(f"PDF saved as {pdf_file}")
except:
    print("Error converting Word file to PDF")