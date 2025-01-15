from re import search
<<<<<<< HEAD
from docx import docx
from docx import Document, opendocx

# from docx import opendocx


class WordUtilities:

=======

import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import os

class WordUtilities:
>>>>>>> main
    def __init__(self):
        pass

    def router(self, cfg):
        if cfg.task == "search_strings":
            self.search_strings(cfg)
<<<<<<< HEAD
        else:
            raise ValueError(f"Task not found in WorkUtilities")
=======
        elif cfg.task == "csv_to_docx":
            self.csv_to_docx(cfg)
        else:
            raise ValueError("Task not found in WorkUtilities")
>>>>>>> main

    def search_strings(self, cfg):
        document_name = cfg["files"]["file_name"]
        document_data = Document(document_name)
        search_string = None
        for paragraph in document_data.paragraphs:
            if search_string in paragraph.text:
                print(True)

        # Open the .docx file
<<<<<<< HEAD
        document = opendocx("A document.docx")

        # Search returns true if found
        search(document, "your search string")
=======
        #document = opendocx("A document.docx")

        # Search returns true if found

    def csv_to_docx(self, cfg):
        
        input_csv = cfg['files']['file_name']

        df = pd.read_csv(input_csv)

        for col in df.select_dtypes(include=['float']).columns:
            df[col] = df[col].round(2)
            
        doc = Document()

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = column_name
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        output_docx = input_csv.replace(".csv", ".docx")
        doc.save(output_docx)
        
        print(f"CSV file has been converted to DOCX file ")
>>>>>>> main
