from typing import Dict, Any
import os
from datetime import datetime
from pathlib import Path
import logging
from .writers import ReportWriter, ReportWriterRegistry
from .dom import DocumentDOM
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""
backlog:
- [ ] #todo doris_writers - redo first section rendering. has to be logo + Title
- [ ] #todo doris_writers - sections are being marked ## instead of #
"""

class DorisMarkdownWriter(ReportWriter):
    """Doris-specific implementation of markdown report writer"""
    
    def __init__(self, dom: DocumentDOM):
        if not isinstance(dom, DocumentDOM):
            raise TypeError("DorisMarkdownWriter expects a DocumentDOM instance")
        super().__init__(dom)
        
        logging.debug(f"Available report details: {self.dom.report_details}")
        
    def flush(self) -> None:
        """Generate markdown report using target_file, ignoring DOM's target_md_file"""
        output_path = Path(os.path.join(self.output_dir, self.target_file))
        logging.info(f"Writing Doris markdown report to: {output_path}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write first section
            f.write(f"# {self.dom.first_section.name}\n\n")
            for para in self.dom.first_section.paragraphs:
                f.write(f"{para['content']}\n")
            f.write("\n")
 
            # Write remaining sections
            for section in self.dom.sections:
                f.write(f"## {section.name}\n\n")
                
                # Write paragraphs
                for para in section.paragraphs:
                    if para["type"] == "markdown_table":
                        f.write(f"{para['content']}\n\n")
                    else:  # text type
                        f.write(f"{para['content']}\n\n")
                
                # Write subsections if any
                if hasattr(section, 'subsections'):
                    for subsection in section.subsections:
                        f.write(f"### {subsection['name']}\n\n")
                        for para in subsection['paragraphs']:
                            if para["type"] == "markdown_table":
                                f.write(f"{para['content']}\n\n")
                            else:  # text type
                                f.write(f"{para['content']}\n\n")

        logging.info(f"Generated Doris markdown report at: {output_path}")

class DorisDocxWriter(ReportWriter):
    """Doris-specific implementation of Word document report writer"""
    
    def __init__(self, dom: DocumentDOM):
        if not isinstance(dom, DocumentDOM):
            raise TypeError("DorisDocxWriter expects a DocumentDOM instance")
        super().__init__(dom)
    
    def _parse_markdown_table(self, markdown_table: str) -> tuple[list[str], list[list[str]]]:
        """Parse markdown table into headers and data rows"""
        lines = markdown_table.strip().split('\n')
        headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
        data_rows = []
        for line in lines[2:]:
            if line.strip():
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                data_rows.append(row)
        return headers, data_rows

    def _add_table_to_doc(self, doc, markdown_table):
        """Add a table to the Word document"""
        headers, data_rows = self._parse_markdown_table(markdown_table)
        
        table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i, row_data in enumerate(data_rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data

    def flush(self) -> None:
        """Generate Word document report"""
        output_path = Path(os.path.join(self.output_dir, self.target_file))
        logging.info(f"Writing Doris Word report to: {output_path}")
        
        doc = Document()
        
        # Write first section
        doc.add_heading(self.dom.first_section.name, level=0)
        for para in self.dom.first_section.paragraphs:
            if para["type"] == "markdown_table":
                self._add_table_to_doc(doc, para['content'])
            else:
                doc.add_paragraph(para['content'])
        
        # Write remaining sections
        for section in self.dom.sections:
            doc.add_heading(section.name, level=1)
            
            # Write paragraphs
            for para in section.paragraphs:
                if para["type"] == "markdown_table":
                    self._add_table_to_doc(doc, para['content'])
                else:
                    doc.add_paragraph(para['content'])
            
            # Write subsections if any
            if hasattr(section, 'subsections'):
                for subsection in section.subsections:
                    doc.add_heading(subsection['name'], level=2)
                    for para in subsection['paragraphs']:
                        if para["type"] == "markdown_table":
                            self._add_table_to_doc(doc, para['content'])
                        else:
                            doc.add_paragraph(para['content'])
        
        # Save document
        doc.save(str(output_path))
        logging.info(f"Generated Doris Word report at: {output_path}")

# Register the writers
ReportWriterRegistry.register("DorisMarkdownWriter", DorisMarkdownWriter)
ReportWriterRegistry.register("DorisDocxWriter", DorisDocxWriter)
