from abc import ABC, abstractmethod
from typing import Optional, Any
from pathlib import Path
import os
import logging
from datetime import datetime
from docx import Document

class ReportWriterRegistry:
    """Registry for report writer implementations"""
    _writers = {}

    @classmethod
    def register(cls, writer_name: str, writer_class: type) -> None:
        """Register a writer implementation"""
        if not issubclass(writer_class, ReportWriter):
            raise TypeError(f"{writer_class.__name__} must be a subclass of ReportWriter")
        cls._writers[writer_name] = writer_class
        logging.info(f"Registered writer: {writer_name}")

    @classmethod
    def get_writer(cls, writer_name: str) -> Optional[type]:
        """Get a writer implementation by name"""
        return cls._writers.get(writer_name)

class ReportWriter(ABC):
    """Abstract base class for report writers"""
    
    def __init__(self, dom: Any):
        self.dom = dom
        self.validate_paths()
        self.target_file = self._generate_target_filename()

    def validate_paths(self) -> None:
        """Validate and normalize paths from DOM"""
        root = os.path.normpath(self.dom.important_paths['project_root'])
        outputs = self.dom.important_paths['outputs_path']
        self.output_dir = Path(os.path.join(root, outputs))
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @abstractmethod
    def flush(self) -> None:
        """Write report to file"""
        pass

    @classmethod
    def validate_writer_class(cls, class_name: str) -> bool:
        """
        Validate if the specified writer class exists and is properly implemented
        """
        writer_class = ReportWriterRegistry.get_writer(class_name)
        if (writer_class and issubclass(writer_class, ReportWriter)):
            return True
        logging.error(f"Writer class {class_name} not found in registry")
        return False

    def _generate_target_filename(self) -> str:
        """Generate standardized output filename"""
        # Get file stub with default fallback
        file_stub = self.dom.report_details.get('target_file_stub', 'default')
        
        # Add timestamp
        timestamp = datetime.now().strftime("%y%m%d-%H%M%S")
        
        # Get extension from report_details
        extension = self.dom.report_details['target_report_type'].lower()
        if extension not in ['md', 'docx']:
            raise ValueError(f"Unsupported report type: {extension}")
        
        return f"{file_stub}.{timestamp}.out.{extension}"

class DefaultMDWriter(ReportWriter):
    """Default implementation of markdown report writer"""
    
    def __init__(self, dom: Any):
        super().__init__(dom)
        
    def flush(self) -> None:
        """Generate markdown report"""
        output_path = Path(os.path.join(self.output_dir, self.target_file))
        logging.info(f"Writing default markdown report to: {output_path}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write first section
            f.write(f"# {self.dom.first_section.name}\n")
            for para in self.dom.first_section.paragraphs:
                f.write(f"{para['content']}\n")
            f.write("\n")

            # Write remaining sections
            for section in self.dom.sections:
                f.write(f"## {section.name}\n")
                
                # Write paragraphs
                for para in section.paragraphs:
                    if para["type"] == "markdown_table":
                        f.write(f"{para['content']}\n")
                    else:  # text type
                        f.write(f"{para['content']}\n")
                
                # Write subsections if any
                if hasattr(section, 'subsections'):
                    for subsection in section.subsections:
                        f.write(f"### {subsection['name']}\n")
                        for para in subsection['paragraphs']:
                            if para["type"] == "markdown_table":
                                f.write(f"{para['content']}\n")
                            else:  # text type
                                f.write(f"{para['content']}\n")

        logging.info(f"Generated default markdown report at: {output_path}")

class DefaultDOCXWriter(ReportWriter):
    """Default implementation of Word document report writer"""
    
    def __init__(self, dom: Any):
        super().__init__(dom)
    
    def _parse_markdown_table(self, markdown_table: str) -> tuple[list[str], list[list[str]]]:
        """Parse markdown table into headers and data rows"""
        lines = markdown_table.strip().split('\n')
        headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
        data_rows = []
        for line in lines[2:]:
            if line.strip():
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                data_rows.append(row)
        return headers, data_rows

    def _add_table_to_doc(self, doc, markdown_table):
        """Add a table to the Word document"""
        headers, data_rows = self._parse_markdown_table(markdown_table)
        
        table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i, row_data in enumerate(data_rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data

    def flush(self) -> None:
        """Generate Word document report"""
        output_path = Path(os.path.join(self.output_dir, self.target_file))
        logging.info(f"Writing default Word report to: {output_path}")
        
        doc = Document()
        
        # Write first section
        doc.add_heading(self.dom.first_section.name, level=0)
        for para in self.dom.first_section.paragraphs:
            if para["type"] == "markdown_table":
                self._add_table_to_doc(doc, para['content'])
            else:
                doc.add_paragraph(para['content'])
        
        # Write remaining sections
        for section in self.dom.sections:
            doc.add_heading(section.name, level=1)
            
            # Write paragraphs
            for para in section.paragraphs:
                if para["type"] == "markdown_table":
                    self._add_table_to_doc(doc, para['content'])
                else:
                    doc.add_paragraph(para['content'])
            
            # Write subsections if any
            if hasattr(section, 'subsections'):
                for subsection in section.subsections:
                    doc.add_heading(subsection['name'], level=2)
                    for para in subsection['paragraphs']:
                        if para["type"] == "markdown_table":
                            self._add_table_to_doc(doc, para['content'])
                        else:
                            doc.add_paragraph(para['content'])
        
        # Save document
        doc.save(str(output_path))
        logging.info(f"Generated default Word report at: {output_path}")

# Register the writers
ReportWriterRegistry.register("DefaultMDWriter", DefaultMDWriter)
ReportWriterRegistry.register("DefaultDOCXWriter", DefaultDOCXWriter)

# - [ ] #todo refactor this 
#   - [ ] centralize file path generation and validation (? or want to move to reportgen level probably not.. but just dropping the idea here) 
#   - [ ] ensure both classes stick to output file naming conventions followed by DorisMarkdownWriter and DorisDocxWriter classes 
#   - [ ] clean up and refactor templates/modules/reportgen/reportgen-cfg.yml removes un-necessary elements like 
#         target_md_file, target_word_file. 
#         use conventions followed by reportgen-cfg-20in-docx.yml :   target_file_stub, report_writer_class, target_report_type
#   - [ ] pull out _flush_markdown from reportgen.py into a new class here called DefaultMarkdownWriter 
#   - [ ] pull out _flush_docx from reportgen.py into a new class here called DefaultDocxWriter
#   - [ ] register these new classes in ReportWriterRegistry
#   - [ ] update reportgen.py to use these new classes
#   - [ ] clean up and refactor reportgen.py to remove the old _flush_markdown and _flush_docx methods
